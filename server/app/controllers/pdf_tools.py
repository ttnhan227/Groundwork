import uuid
import zipfile
from io import BytesIO

import fitz
from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Response,
    UploadFile,
    status,
)
from fastapi.responses import StreamingResponse
from PIL import Image, ImageDraw, ImageOps
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.concurrency import run_in_threadpool

from app.database import get_session
from app.dependencies import current_user
from app.documents import owned_document, safe_filename
from app.models import (
    ArtifactVersion,
    Document,
    DocumentStatus,
    GeneratedArtifact,
    JobStatus,
    ProcessingJob,
    User,
)
from app.pdf_operations import (
    delete_pages,
    images_to_pdf,
    merge_pdfs,
    pdf_to_images,
    rotate_pages,
    select_pages,
    split_pdf,
    watermark_pdf,
)
from app.schemas import (
    ArtifactRenameRequest,
    ArtifactResponse,
    ArtifactVersionResponse,
    ArtifactVersionRestoreRequest,
    DocumentResponse,
    MergeRequest,
    PageOperationRequest,
    PDFToImagesRequest,
    RotateRequest,
    SplitRequest,
)
from app.storage import ObjectStorage

router = APIRouter(prefix="/pdf-tools", tags=["PDF tools"])


def _render_pdf_preview(data: bytes) -> bytes:
    pdf = fitz.open(stream=data, filetype="pdf")
    try:
        if not pdf.page_count:
            raise ValueError("PDF has no pages")
        return pdf[0].get_pixmap(matrix=fitz.Matrix(.7, .7), alpha=False).tobytes("png")
    finally:
        pdf.close()


def _render_image_preview(data: bytes) -> bytes:
    with Image.open(BytesIO(data)) as source:
        image = ImageOps.exif_transpose(source).convert("RGB")
        image.thumbnail((420, 540))
        canvas = Image.new("RGB", (420, 540), "white")
        x = (canvas.width - image.width) // 2
        y = (canvas.height - image.height) // 2
        canvas.paste(image, (x, y))
        output = BytesIO()
        canvas.save(output, "PNG", optimize=True)
        return output.getvalue()


def _render_file_card(filename: str, content_type: str, excerpt: str = "") -> bytes:
    image = Image.new("RGB", (420, 540), "#f7f5fb")
    draw = ImageDraw.Draw(image)
    extension = (filename.rsplit(".", 1)[-1] if "." in filename else "FILE").upper()[:8]
    draw.rounded_rectangle((28, 28, 392, 512), radius=22, fill="white", outline="#ddd7e8", width=3)
    draw.rounded_rectangle((55, 58, 180, 112), radius=12, fill="#6f5bd3")
    draw.text((78, 75), extension, fill="white")
    display_name = filename[:34] + ("…" if len(filename) > 34 else "")
    draw.text((55, 145), display_name, fill="#252132")
    draw.text((55, 178), content_type[:45], fill="#777184")
    lines = [line.strip() for line in excerpt.replace("\r", "").split("\n") if line.strip()]
    y = 225
    for line in lines[:8]:
        wrapped = [line[index:index + 45] for index in range(0, min(len(line), 135), 45)]
        for part in wrapped:
            draw.text((55, y), part, fill="#514b5d")
            y += 22
            if y > 455:
                break
        if y > 455:
            break
    if not lines:
        draw.text((55, 235), "Ready to download", fill="#777184")
    output = BytesIO()
    image.save(output, "PNG", optimize=True)
    return output.getvalue()


def _artifact_preview(artifact: GeneratedArtifact) -> bytes:
    data = ObjectStorage().download(artifact.object_key)
    content_type = artifact.content_type.lower()
    filename = artifact.filename.lower()
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        return _render_pdf_preview(data)
    if content_type.startswith("image/") or filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return _render_image_preview(data)
    if content_type == "application/zip" or filename.endswith(".zip"):
        with zipfile.ZipFile(BytesIO(data)) as archive:
            for member in archive.infolist():
                if member.is_dir() or member.file_size > 20 * 1024 * 1024:
                    continue
                member_name = member.filename.lower()
                member_data = archive.read(member)
                if member_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
                    return _render_image_preview(member_data)
                if member_name.endswith(".pdf"):
                    return _render_pdf_preview(member_data)
            names = "\n".join(member.filename for member in archive.infolist()[:8])
            return _render_file_card(artifact.filename, artifact.content_type, names)
    excerpt = ""
    if filename.endswith((".md", ".txt")) or content_type.startswith("text/"):
        excerpt = data[:4000].decode("utf-8", errors="replace")
    elif filename.endswith(".docx"):
        from docx import Document as WordDocument
        word = WordDocument(BytesIO(data))
        excerpt = "\n".join(paragraph.text for paragraph in word.paragraphs[:10])
    return _render_file_card(artifact.filename, artifact.content_type, excerpt)


def _artifact_as_pdf(artifact: GeneratedArtifact) -> bytes:
    data = ObjectStorage().download(artifact.object_key)
    content_type = artifact.content_type.lower()
    filename = artifact.filename.lower()
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        return data
    if content_type.startswith("image/") or filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return images_to_pdf([data])
    if content_type == "application/zip" or filename.endswith(".zip"):
        parts: list[bytes] = []
        with zipfile.ZipFile(BytesIO(data)) as archive:
            for member in archive.infolist():
                if member.is_dir() or member.file_size > 20 * 1024 * 1024:
                    continue
                member_data = archive.read(member)
                member_name = member.filename.lower()
                if member_name.endswith(".pdf"):
                    parts.append(member_data)
                elif member_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
                    parts.append(images_to_pdf([member_data]))
        if not parts:
            raise ValueError("This ZIP does not contain PDFs or supported images")
        return merge_pdfs(parts)
    text = ""
    if filename.endswith((".md", ".txt")) or content_type.startswith("text/"):
        text = data.decode("utf-8", errors="replace")
    elif filename.endswith(".docx"):
        from docx import Document as WordDocument
        word = WordDocument(BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    if text.strip():
        pdf = fitz.open()
        try:
            lines = text.splitlines() or [text]
            for start in range(0, len(lines), 45):
                page = pdf.new_page()
                page.insert_textbox(
                    fitz.Rect(54, 54, page.rect.width - 54, page.rect.height - 54),
                    "\n".join(lines[start:start + 45]),
                    fontsize=10,
                )
            return pdf.tobytes(garbage=4, deflate=True)
        finally:
            pdf.close()
    raise ValueError("This file type cannot be indexed for PDF or AI tools")


async def _owned_artifact(
    artifact_id: uuid.UUID, user: User, session: AsyncSession
) -> GeneratedArtifact:
    artifact = await session.scalar(
        select(GeneratedArtifact).where(
            GeneratedArtifact.id == artifact_id,
            GeneratedArtifact.owner_id == user.id,
        )
    )
    if artifact is None:
        raise HTTPException(status_code=404, detail="Generated file not found")
    return artifact


async def _ready_document(identifier: uuid.UUID, user: User, session: AsyncSession) -> Document:
    document = await owned_document(identifier, user, session)
    if document.status != DocumentStatus.READY:
        raise HTTPException(status_code=409, detail="The document must finish processing first")
    return document


async def _store(
    operation: str,
    filename: str,
    data: bytes,
    content_type: str,
    parameters: dict,
    user: User,
    session: AsyncSession,
) -> GeneratedArtifact:
    identifier = uuid.uuid4()
    filename = safe_filename(filename)
    key = f"{user.id}/generated/{identifier}/{filename}"
    try:
        ObjectStorage().upload(key, data, content_type)
    except Exception as exc:
        raise HTTPException(status_code=503, detail="Generated-file storage is temporarily unavailable") from exc
    from app.deliverables import ensure_personal_workspace
    workspace = await ensure_personal_workspace(user, session)
    artifact = GeneratedArtifact(
        id=identifier, owner_id=user.id, workspace_id=workspace.id, operation=operation, filename=filename,
        object_key=key, content_type=content_type, size_bytes=len(data), parameters=parameters,
    )
    session.add(artifact)
    session.add(ArtifactVersion(
        artifact_id=identifier,
        version_number=1,
        object_key=key,
        content_type=content_type,
        size_bytes=len(data),
        metadata_json={"operation": operation},
    ))
    await session.commit()
    await session.refresh(artifact)
    return artifact


def _run(operation):
    try:
        return operation()
    except (ValueError, TypeError, RuntimeError) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc


@router.get("/artifacts", response_model=list[ArtifactResponse])
async def list_artifacts(
    user: User = Depends(current_user), session: AsyncSession = Depends(get_session)
) -> list[GeneratedArtifact]:
    return list(await session.scalars(
        select(GeneratedArtifact).where(GeneratedArtifact.owner_id == user.id).order_by(GeneratedArtifact.created_at.desc())
    ))


@router.get("/artifacts/{artifact_id}/thumbnail")
async def artifact_thumbnail(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> Response:
    artifact = await _owned_artifact(artifact_id, user, session)
    try:
        image = await run_in_threadpool(_artifact_preview, artifact)
    except Exception:
        image = _render_file_card(artifact.filename, artifact.content_type)
    return Response(
        content=image,
        media_type="image/png",
        headers={"Cache-Control": "private, max-age=3600"},
    )


@router.get("/artifacts/{artifact_id}/download")
async def download_artifact(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> StreamingResponse:
    artifact = await _owned_artifact(artifact_id, user, session)
    data = ObjectStorage().download(artifact.object_key)
    return StreamingResponse(
        BytesIO(data), media_type=artifact.content_type,
        headers={"Content-Disposition": f'attachment; filename="{safe_filename(artifact.filename)}"'},
    )


@router.get("/artifacts/{artifact_id}", response_model=ArtifactResponse)
async def get_artifact(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    return await _owned_artifact(artifact_id, user, session)


@router.get(
    "/artifacts/{artifact_id}/versions",
    response_model=list[ArtifactVersionResponse],
)
async def list_artifact_versions(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> list[ArtifactVersion]:
    await _owned_artifact(artifact_id, user, session)
    return list(await session.scalars(
        select(ArtifactVersion)
        .where(ArtifactVersion.artifact_id == artifact_id)
        .order_by(ArtifactVersion.version_number.desc())
    ))


@router.post(
    "/artifacts/{artifact_id}/versions/restore",
    response_model=ArtifactVersionResponse,
)
async def restore_artifact_version(
    artifact_id: uuid.UUID,
    payload: ArtifactVersionRestoreRequest,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> ArtifactVersion:
    artifact = await _owned_artifact(artifact_id, user, session)
    source = await session.scalar(
        select(ArtifactVersion).where(
            ArtifactVersion.id == payload.version_id,
            ArtifactVersion.artifact_id == artifact_id,
        )
    )
    if source is None:
        raise HTTPException(status_code=404, detail="Artifact version not found")
    latest = await session.scalar(
        select(func.max(ArtifactVersion.version_number)).where(
            ArtifactVersion.artifact_id == artifact_id
        )
    )
    restored = ArtifactVersion(
        artifact_id=artifact_id,
        version_number=(latest or 0) + 1,
        object_key=source.object_key,
        content_type=source.content_type,
        size_bytes=source.size_bytes,
        change_prompt=f"Restored version {source.version_number}",
        metadata_json={
            **source.metadata_json,
            "restored_from_version": source.version_number,
        },
    )
    artifact.object_key = source.object_key
    artifact.content_type = source.content_type
    artifact.size_bytes = source.size_bytes
    session.add(restored)
    await session.commit()
    await session.refresh(restored)
    return restored


@router.post("/artifacts/{artifact_id}/duplicate", response_model=ArtifactResponse, status_code=201)
async def duplicate_artifact(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    artifact = await _owned_artifact(artifact_id, user, session)
    data = ObjectStorage().download(artifact.object_key)
    stem, dot, extension = artifact.filename.rpartition(".")
    filename = f"{stem or artifact.filename}-copy{dot}{extension}" if dot else f"{artifact.filename}-copy"
    return await _store(
        "duplicate",
        filename,
        data,
        artifact.content_type,
        {**artifact.parameters, "duplicated_from": str(artifact.id)},
        user,
        session,
    )


@router.post("/artifacts/{artifact_id}/index", response_model=DocumentResponse)
async def index_artifact(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> Document:
    artifact = await _owned_artifact(artifact_id, user, session)
    if artifact.linked_document_id:
        linked = await session.get(Document, artifact.linked_document_id)
        if linked is not None and linked.owner_id == user.id:
            return linked

    try:
        data = await run_in_threadpool(_artifact_as_pdf, artifact)
    except (ValueError, TypeError, RuntimeError, zipfile.BadZipFile) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    document_id = uuid.uuid4()
    filename = artifact.filename.rsplit(".", 1)[0] + ".pdf"
    object_key = f"{user.id}/{document_id}/{safe_filename(filename)}"
    try:
        ObjectStorage().upload_pdf(object_key, data)
    except Exception as exc:
        raise HTTPException(status_code=503, detail="Document storage is temporarily unavailable") from exc

    document = Document(
        id=document_id,
        owner_id=user.id,
        workspace_id=artifact.workspace_id,
        filename=safe_filename(filename),
        object_key=object_key,
        content_type="application/pdf",
        size_bytes=len(data),
    )
    job = ProcessingJob(
        document_id=document_id,
        owner_id=user.id,
        operation="document_processing",
        status=JobStatus.QUEUED,
        progress=0,
    )
    session.add_all([document, job])
    artifact.linked_document_id = document_id
    await session.commit()
    await session.refresh(document)

    from app.tasks import process_document
    try:
        task = process_document.delay(str(document.id))
        job.task_id = task.id
        await session.commit()
    except Exception as exc:
        document.status = DocumentStatus.FAILED
        document.error_message = "Processing queue is temporarily unavailable"
        job.status = JobStatus.FAILED
        job.error_message = document.error_message
        await session.commit()
        raise HTTPException(status_code=503, detail=document.error_message) from exc
    return document


@router.patch("/artifacts/{artifact_id}", response_model=ArtifactResponse)
async def rename_artifact(
    artifact_id: uuid.UUID,
    payload: ArtifactRenameRequest,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    artifact = await _owned_artifact(artifact_id, user, session)
    artifact.filename = safe_filename(payload.filename)
    if artifact.linked_document_id:
        linked = await session.get(Document, artifact.linked_document_id)
        if linked is not None:
            linked.filename = artifact.filename.rsplit(".", 1)[0] + ".pdf"
    await session.commit()
    await session.refresh(artifact)
    return artifact


@router.delete("/artifacts/{artifact_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_artifact(
    artifact_id: uuid.UUID,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> Response:
    artifact = await _owned_artifact(artifact_id, user, session)
    linked = await session.get(Document, artifact.linked_document_id) if artifact.linked_document_id else None
    try:
        ObjectStorage().remove(artifact.object_key)
        if linked is not None:
            ObjectStorage().remove(linked.object_key)
    except Exception as exc:
        raise HTTPException(status_code=503, detail="Generated-file storage is temporarily unavailable") from exc
    if linked is not None:
        await session.delete(linked)
    await session.delete(artifact)
    await session.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.post("/merge", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def merge(
    payload: MergeRequest,
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    if len(set(payload.document_ids)) != len(payload.document_ids):
        raise HTTPException(status_code=422, detail="Select each PDF only once")
    documents = [await _ready_document(item, user, session) for item in payload.document_ids]
    data = _run(lambda: merge_pdfs([ObjectStorage().download(document.object_key) for document in documents]))
    return await _store("merge", "merged.pdf", data, "application/pdf", payload.model_dump(mode="json"), user, session)


async def _page_operation(
    operation: str, suffix: str, payload: PageOperationRequest, transform, user: User, session: AsyncSession
) -> GeneratedArtifact:
    document = await _ready_document(payload.document_id, user, session)
    data = _run(lambda: transform(ObjectStorage().download(document.object_key)))
    filename = f"{document.filename.removesuffix('.pdf')}-{suffix}.pdf"
    return await _store(operation, filename, data, "application/pdf", payload.model_dump(mode="json"), user, session)


@router.post("/extract", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def extract_pages_route(payload: PageOperationRequest, user: User = Depends(current_user), session: AsyncSession = Depends(get_session)):
    return await _page_operation("extract", "extracted", payload, lambda data: select_pages(data, payload.page_numbers), user, session)


@router.post("/delete-pages", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def delete_pages_route(payload: PageOperationRequest, user: User = Depends(current_user), session: AsyncSession = Depends(get_session)):
    return await _page_operation("delete_pages", "pages-removed", payload, lambda data: delete_pages(data, payload.page_numbers), user, session)


@router.post("/rotate", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def rotate(payload: RotateRequest, user: User = Depends(current_user), session: AsyncSession = Depends(get_session)):
    return await _page_operation("rotate", "rotated", payload, lambda data: rotate_pages(data, payload.page_numbers, payload.degrees), user, session)


@router.post("/split", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def split(payload: SplitRequest, user: User = Depends(current_user), session: AsyncSession = Depends(get_session)):
    document = await _ready_document(payload.document_id, user, session)
    data = _run(lambda: split_pdf(ObjectStorage().download(document.object_key), payload.mode, payload.ranges, payload.page_numbers))
    return await _store("split", f"{document.filename.removesuffix('.pdf')}-split.zip", data, "application/zip", payload.model_dump(mode="json"), user, session)


@router.post("/pdf-to-images", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def convert_pdf_to_images(payload: PDFToImagesRequest, user: User = Depends(current_user), session: AsyncSession = Depends(get_session)):
    document = await _ready_document(payload.document_id, user, session)
    data = _run(lambda: pdf_to_images(ObjectStorage().download(document.object_key), payload.page_numbers, payload.format, payload.dpi))
    return await _store("pdf_to_images", f"{document.filename.removesuffix('.pdf')}-images.zip", data, "application/zip", payload.model_dump(mode="json"), user, session)


@router.post("/images-to-pdf", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def convert_images_to_pdf(
    files: list[UploadFile] = File(...),
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    if not 1 <= len(files) <= 50:
        raise HTTPException(status_code=422, detail="Upload between 1 and 50 images")
    allowed = {"image/png", "image/jpeg"}
    if any(file.content_type not in allowed for file in files):
        raise HTTPException(status_code=415, detail="Only PNG and JPEG images are accepted")
    content = [await file.read(20 * 1024 * 1024 + 1) for file in files]
    if any(len(item) > 20 * 1024 * 1024 for item in content):
        raise HTTPException(status_code=413, detail="Each image must be 20 MB or smaller")
    data = _run(lambda: images_to_pdf(content))
    return await _store("images_to_pdf", "images.pdf", data, "application/pdf", {"image_count": len(files)}, user, session)


@router.post("/watermark", response_model=ArtifactResponse, status_code=status.HTTP_201_CREATED)
async def watermark(
    document_id: uuid.UUID = Form(...),
    text: str | None = Form(default=None, max_length=120),
    page_numbers: str = Form(default=""),
    position: str = Form(default="center"),
    opacity: float = Form(default=.25),
    rotation: int = Form(default=0),
    image: UploadFile | None = File(default=None),
    user: User = Depends(current_user),
    session: AsyncSession = Depends(get_session),
) -> GeneratedArtifact:
    document = await _ready_document(document_id, user, session)
    image_data = None
    if image:
        if image.content_type not in {"image/png", "image/jpeg"}:
            raise HTTPException(status_code=415, detail="Watermark image must be PNG or JPEG")
        image_data = await image.read(10 * 1024 * 1024 + 1)
        if len(image_data) > 10 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="Watermark image must be 10 MB or smaller")
    try:
        pages = [int(item.strip()) for item in page_numbers.split(",") if item.strip()] or None
    except ValueError as exc:
        raise HTTPException(status_code=422, detail="Pages must be comma-separated numbers") from exc
    data = _run(lambda: watermark_pdf(
        ObjectStorage().download(document.object_key), text.strip() if text else None,
        image_data, pages, position, opacity, rotation,
    ))
    parameters = {"text": text, "page_numbers": pages, "position": position, "opacity": opacity, "rotation": rotation, "has_image": bool(image)}
    return await _store("watermark", f"{document.filename.removesuffix('.pdf')}-watermarked.pdf", data, "application/pdf", parameters, user, session)
