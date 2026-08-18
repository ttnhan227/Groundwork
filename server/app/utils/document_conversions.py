import os
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

import fitz
from docx import Document as WordDocument
from docx.document import Document as WordDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


def validate_docx(data: bytes) -> None:
    try:
        with ZipFile(BytesIO(data)) as archive:
            if "word/document.xml" not in archive.namelist():
                raise ValueError("The uploaded file is not a valid DOCX document")
    except BadZipFile as exc:
        raise ValueError("The uploaded file is not a valid DOCX document") from exc


def docx_to_pdf(data: bytes) -> bytes:
    validate_docx(data)
    with tempfile.TemporaryDirectory(prefix="groundwork-docx-") as directory:
        workdir = Path(directory)
        source = workdir / "source.docx"
        source.write_bytes(data)
        environment = os.environ.copy()
        environment["HOME"] = str(workdir)
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(workdir),
                str(source),
            ],
            capture_output=True,
            check=False,
            env=environment,
            timeout=120,
        )
        output = workdir / "source.pdf"
        if result.returncode != 0 or not output.exists():
            detail = result.stderr.decode("utf-8", errors="replace").strip()
            raise RuntimeError(detail or "Word-to-PDF conversion failed")
        converted = output.read_bytes()
    document = fitz.open(stream=converted, filetype="pdf")
    try:
        if document.page_count < 1:
            raise RuntimeError("Word-to-PDF conversion produced an empty document")
    finally:
        document.close()
    return converted


def _iter_blocks(document: WordDocumentType):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _markdown_text(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        value = run.text.replace("\\", "\\\\").replace("*", "\\*").replace("_", "\\_")
        if not value:
            continue
        if run.bold:
            value = f"**{value}**"
        if run.italic:
            value = f"*{value}*"
        parts.append(value)
    return "".join(parts).strip()


def docx_to_markdown(data: bytes) -> bytes:
    validate_docx(data)
    document = WordDocument(BytesIO(data))
    lines: list[str] = []
    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            text = _markdown_text(block)
            if not text:
                if lines and lines[-1] != "":
                    lines.append("")
                continue
            style = (block.style.name if block.style else "").lower()
            if style == "title":
                text = f"# {text}"
            elif style == "subtitle":
                text = f"*{text}*"
            elif style.startswith("heading"):
                try:
                    level = min(6, max(1, int(style.split()[-1])))
                except ValueError:
                    level = 2
                text = f"{'#' * level} {text}"
            elif "list bullet" in style:
                text = f"- {text}"
            elif "list number" in style:
                text = f"1. {text}"
            lines.extend([text, ""])
        else:
            rows = [
                [
                    cell.text.replace("|", "\\|").replace("\n", "<br>")
                    for cell in row.cells
                ]
                for row in block.rows
            ]
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
            lines.append("")
    content = "\n".join(lines).strip() + "\n"
    return content.encode("utf-8")


def pdf_to_docx(data: bytes) -> bytes:
    source = fitz.open(stream=data, filetype="pdf")
    if source.page_count < 1:
        source.close()
        raise ValueError("The PDF has no pages")
    output = WordDocument()
    normal = output.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for page_index, page in enumerate(source):
        if page_index:
            output.add_page_break()
        blocks = page.get_text("dict", sort=True).get("blocks", [])
        for block in blocks:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    paragraph = output.add_paragraph()
                    paragraph.paragraph_format.space_after = Pt(2)
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue
                        run = paragraph.add_run(text)
                        flags = int(span.get("flags", 0))
                        run.bold = bool(flags & 16)
                        run.italic = bool(flags & 2)
                        run.font.size = Pt(min(28, max(8, float(span.get("size", 10)))))
                    if not paragraph.text:
                        paragraph._element.getparent().remove(paragraph._element)
            elif block.get("type") == 1 and block.get("image"):
                try:
                    bbox = block.get("bbox", (0, 0, 360, 240))
                    width = min(6.5, max(1.0, (bbox[2] - bbox[0]) / 72))
                    output.add_picture(BytesIO(block["image"]), width=Inches(width))
                except Exception:
                    continue
    source.close()
    stream = BytesIO()
    output.save(stream)
    converted = stream.getvalue()
    validate_docx(converted)
    return converted
