import shutil
from io import BytesIO

import fitz
import pytest
from docx import Document

from app.document_conversions import (
    docx_to_markdown,
    docx_to_pdf,
    pdf_to_docx,
    validate_docx,
)


def sample_docx() -> bytes:
    document = Document()
    document.add_heading("Quarterly Report", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Important").bold = True
    paragraph.add_run(" project update")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Taylor"
    table.cell(1, 1).text = "Complete"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def sample_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Quarterly Report")
    page.insert_text((72, 100), "Editable PDF content")
    data = document.tobytes()
    document.close()
    return data


def test_docx_to_markdown_preserves_structure() -> None:
    markdown = docx_to_markdown(sample_docx()).decode()
    assert "# Quarterly Report" in markdown
    assert "**Important** project update" in markdown
    assert "| Owner | Status |" in markdown
    assert "| Taylor | Complete |" in markdown


def test_pdf_to_docx_creates_editable_word_document() -> None:
    converted = pdf_to_docx(sample_pdf())
    validate_docx(converted)
    document = Document(BytesIO(converted))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Quarterly Report" in text
    assert "Editable PDF content" in text


@pytest.mark.skipif(
    shutil.which("soffice") is None,
    reason="LibreOffice is not installed",
)
def test_docx_to_pdf_creates_valid_pdf() -> None:
    converted = docx_to_pdf(sample_docx())
    document = fitz.open(stream=converted, filetype="pdf")
    assert document.page_count == 1
    assert "Quarterly Report" in document[0].get_text()
    document.close()


def test_invalid_docx_is_rejected() -> None:
    with pytest.raises(ValueError, match="valid DOCX"):
        validate_docx(b"not a Word document")
